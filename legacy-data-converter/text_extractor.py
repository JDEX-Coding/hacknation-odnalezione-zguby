"""
Text extraction module for various file formats.
Supports: PDF, DOCX, HTML, TXT, JSON, XML, CSV
"""

import os
import json
import csv
from typing import Dict, List, Optional
from pathlib import Path
import logging

# PDF extraction
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# DOCX extraction
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# HTML extraction
try:
    from bs4 import BeautifulSoup
    HTML_AVAILABLE = True
except ImportError:
    HTML_AVAILABLE = False

# XML parsing
import xml.etree.ElementTree as ET

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class TextExtractor:
    """Extract text from various file formats."""

    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.html', '.htm', '.txt', '.json', '.xml', '.csv']

    def extract(self, file_path: str) -> Dict[str, any]:
        """
        Extract text from a file.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Dict containing:
                - text: Extracted text content
                - metadata: File metadata (format, size, etc.)
                - raw_data: Original structured data (for JSON, CSV, XML)
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = file_path.suffix.lower()
        
        if extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {extension}")
        
        logger.info(f"Extracting text from: {file_path.name}")
        
        result = {
            "text": "",
            "metadata": {
                "file_name": file_path.name,
                "file_format": extension,
                "file_size": file_path.stat().st_size
            },
            "raw_data": None
        }
        
        if extension == '.pdf':
            result["text"] = self._extract_pdf(file_path)
        elif extension == '.docx':
            result["text"] = self._extract_docx(file_path)
        elif extension in ['.html', '.htm']:
            result["text"] = self._extract_html(file_path)
        elif extension == '.txt':
            result["text"] = self._extract_txt(file_path)
        elif extension == '.json':
            text, raw = self._extract_json(file_path)
            result["text"] = text
            result["raw_data"] = raw
        elif extension == '.xml':
            text, raw = self._extract_xml(file_path)
            result["text"] = text
            result["raw_data"] = raw
        elif extension == '.csv':
            text, raw = self._extract_csv(file_path)
            result["text"] = text
            result["raw_data"] = raw
        
        return result

    def _extract_pdf(self, file_path: Path) -> str:
        """Extract text from PDF."""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 not installed. Run: pip install PyPDF2")
        
        text_parts = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    text = page.extract_text()
                    if text:
                        text_parts.append(f"[Page {page_num + 1}]\n{text}")
                except Exception as e:
                    logger.warning(f"Error extracting page {page_num + 1}: {e}")
        
        return "\n\n".join(text_parts)

    def _extract_docx(self, file_path: Path) -> str:
        """Extract text from DOCX."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        
        doc = Document(file_path)
        text_parts = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)
        
        return "\n".join(text_parts)

    def _extract_html(self, file_path: Path) -> str:
        """Extract text from HTML."""
        if not HTML_AVAILABLE:
            raise ImportError("beautifulsoup4 not installed. Run: pip install beautifulsoup4")
        
        with open(file_path, 'r', encoding='utf-8') as file:
            soup = BeautifulSoup(file.read(), 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            return text

    def _extract_txt(self, file_path: Path) -> str:
        """Extract text from plain text file."""
        encodings = ['utf-8', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    return file.read()
            except UnicodeDecodeError:
                continue
        
        raise ValueError(f"Could not decode text file with supported encodings")

    def _extract_json(self, file_path: Path) -> tuple[str, any]:
        """Extract text from JSON."""
        with open(file_path, 'r', encoding='utf-8') as file:
            data = json.load(file)
        
        # Convert JSON to readable text
        text = json.dumps(data, indent=2, ensure_ascii=False)
        
        return text, data

    def _extract_xml(self, file_path: Path) -> tuple[str, any]:
        """Extract text from XML."""
        tree = ET.parse(file_path)
        root = tree.getroot()
        
        # Extract all text content
        text_parts = []
        
        def traverse(element, level=0):
            indent = "  " * level
            if element.text and element.text.strip():
                text_parts.append(f"{indent}{element.tag}: {element.text.strip()}")
            
            for child in element:
                traverse(child, level + 1)
        
        traverse(root)
        text = "\n".join(text_parts)
        
        # Convert to dict for structured access
        def elem_to_dict(element):
            result = {}
            if element.text and element.text.strip():
                result['_text'] = element.text.strip()
            if element.attrib:
                result['_attributes'] = element.attrib
            for child in element:
                child_data = elem_to_dict(child)
                if child.tag in result:
                    if not isinstance(result[child.tag], list):
                        result[child.tag] = [result[child.tag]]
                    result[child.tag].append(child_data)
                else:
                    result[child.tag] = child_data
            return result
        
        raw_data = {root.tag: elem_to_dict(root)}
        
        return text, raw_data

    def _extract_csv(self, file_path: Path) -> tuple[str, List[Dict]]:
        """Extract text from CSV."""
        rows = []
        text_parts = []
        
        encodings = ['utf-8', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding, newline='') as file:
                    # Try to detect delimiter
                    sample = file.read(1024)
                    file.seek(0)
                    sniffer = csv.Sniffer()
                    try:
                        delimiter = sniffer.sniff(sample).delimiter
                    except:
                        delimiter = ','
                    
                    reader = csv.DictReader(file, delimiter=delimiter)
                    
                    # Add header
                    if reader.fieldnames:
                        text_parts.append(" | ".join(reader.fieldnames))
                        text_parts.append("-" * 80)
                    
                    for row in reader:
                        rows.append(row)
                        text_parts.append(" | ".join(str(v) for v in row.values()))
                    
                    break
            except UnicodeDecodeError:
                continue
        
        if not rows:
            raise ValueError(f"Could not read CSV file with supported encodings")
        
        text = "\n".join(text_parts)
        
        return text, rows


def extract_from_file(file_path: str) -> Dict[str, any]:
    """
    Convenience function to extract text from a file.
    
    Args:
        file_path: Path to the file
        
    Returns:
        Dict with text, metadata, and raw_data
    """
    extractor = TextExtractor()
    return extractor.extract(file_path)


if __name__ == "__main__":
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python text_extractor.py <file_path>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    
    try:
        result = extract_from_file(file_path)
        print(f"\n{'='*80}")
        print(f"File: {result['metadata']['file_name']}")
        print(f"Format: {result['metadata']['file_format']}")
        print(f"Size: {result['metadata']['file_size']} bytes")
        print(f"{'='*80}\n")
        print(result['text'][:1000])  # Print first 1000 chars
        if len(result['text']) > 1000:
            print(f"\n... ({len(result['text']) - 1000} more characters)")
    except Exception as e:
        logger.error(f"Error: {e}")
        sys.exit(1)
